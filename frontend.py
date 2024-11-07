import streamlit as st
import requests
import time
import pyperclip  # For clipboard operations
import tempfile  # To create temporary files
from docx import Document  # Importing Document class from python-docx

# Set the page layout and title with a more modern theme
st.set_page_config(
    page_title="Assignment Creator",
    page_icon="📘",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Add a custom CSS for enhanced styling
st.markdown("""
    <style>
    body {
        background-color: #f4f6f9;
        font-family: 'Roboto', sans-serif;
    }
    h1 {
        text-align: center;
        background: linear-gradient(to right, #8e2de2, #4a00e0);
        -webkit-background-clip: text;
        color: transparent;
        font-size: 2.5em;
        margin-bottom: 20px;
    }
    .stTextInput input {
        border: 2px solid #4a00e0;
        padding: 10px;
        border-radius: 10px;
    }
    .stSelectbox .st-bt, .stNumberInput input {
        border: 2px solid #8e2de2;
        border-radius: 10px;
    }
    .stButton button {
        background-color: #8e2de2 !important;
        border-radius: 10px !important;
        font-size: 16px !important;
    }
    .stButton button:hover {
        background-color: #6c24c5 !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Add a title with a gradient effect centered
st.markdown("""
    <h2 style="text-align: center; font-size: 2.5em;">
        Let's Create An Assignment for You.
    </h2>
    """, unsafe_allow_html=True)

# Add an animated image/gif while explaining the process
# st.markdown("<center><img src='https://i.giphy.com/media/v1.Y2lkPTc5MGI3NjExa3doaGx6dW15dWlpdnZuNDJjOGdyZTR1NjVyZTV2YXNhN3plNDZqbiZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/kZhvGXpT7bOsltxydL/giphy.gif' width='400'></center>", unsafe_allow_html=True)

# Initialize session state for generated content to prevent loss during reruns
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = None  # Initialize as None

# Form to capture input details, using columns for a structured layout
with st.form(key="user_form"):
    col1, col2 = st.columns(2)
    with col1:
        institution_name = st.text_input("Institution Name", placeholder="🏫 Institution Name", help="Enter the name of your institution")
        name = st.text_input("Your Name", placeholder="✍️ John Doe", help="Enter your full name")
        roll_no = st.text_input("Roll No.", placeholder="🆔 FA22-BSCS-342", help="Enter your roll number")
    with col2:
        course_name = st.text_input("Course Name", placeholder="🎓 Course Name", help="Enter the course you're taking")
        instructor_name = st.text_input("Instructor Name", placeholder="👩‍🏫 Instructor Name", help="Enter your instructor's name")
        assignment_topic = st.text_input("Assignment Topic", placeholder="📝 Assignment Topic", help="Enter the assignment topic")

    st.markdown("<hr>", unsafe_allow_html=True)  # Add a horizontal divider

    # Input for headings and word count
    heading_no = st.number_input("Number of Headings", min_value=1, max_value=120, step=1, help="Choose the number of headings")
    assignment_len = st.number_input("Assignment Length (in words)", min_value=100, step=50, help="Choose the word count for your assignment")
    assignment_format = st.selectbox("Select Assignment Format", options=["Formal 📄", "Professional 🏢"], help="Select the style of the assignment")

    # Submit button for generating the assignment
    submit_button = st.form_submit_button(label="📥 Generate Assignment")

# Process and display the results upon submission
if submit_button:
    # Validate all fields
    if not institution_name or not name or not roll_no or not course_name or not instructor_name or not assignment_topic:
        st.error("⚠️ All fields are required. Please fill out the entire form.")
    else:
        backend_url = "http://localhost:8000/generate"  # Update this to match your backend URL

        data = {
            "institution_name": institution_name,
            "name": name,
            "roll_no": roll_no,
            "course_name": course_name,
            "instructor_name": instructor_name,
            "assignment_topic": assignment_topic,
            "heading_no": heading_no,
            "assignment_len": assignment_len,
            "assignment_format": assignment_format
        }

        # Show progress during content generation
        with st.spinner('🚀 Generating your assignment...'):
            progress_bar = st.progress(0)
            for percent_complete in range(100):
                time.sleep(0.05)
                progress_bar.progress(percent_complete + 1)
            
            try:
                # Call the backend API to generate assignment content
                response = requests.post(backend_url, json=data)

                # Check for successful response
                if response.status_code == 200:
                    st.session_state.generated_content = response.json().get("assignment_content", "No content generated.")
                else:
                    st.error(f"❌ Failed to generate the assignment. Error: {response.status_code}")
            except Exception as e:
                st.error(f"❌ An error occurred: {e}")

# If content is generated, display it
if st.session_state.generated_content:
    st.subheader("🎉 Your Generated Assignment:")

    def typing_effect(text, delay=0.05):
        for line in text.split("\n"):
            if line.strip():
                st.markdown(line.strip())  # Display each line properly
                time.sleep(delay)

    typing_effect(st.session_state.generated_content)

    st.markdown("<hr>", unsafe_allow_html=True)  # Add a horizontal divider

    # Copy to clipboard button
    if st.button('📋 Copy to Clipboard'):
        pyperclip.copy(st.session_state.generated_content)
        st.success("✔️ Assignment content copied to clipboard!")

    # DOCX download button
    if st.button('📄 Download as DOCX'):
        try:
            doc = Document()
            doc.add_heading('Generated Assignment', level=1)

            for line in st.session_state.generated_content.split('\n'):
                if line.strip():
                    if line.startswith("###"):
                        doc.add_heading(line[4:], level=2).bold = True
                    else:
                        doc.add_paragraph(line)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx_file:
                doc.save(tmp_docx_file.name)
                docx_file_path = tmp_docx_file.name

            with open(docx_file_path, 'rb') as f:
                docx_bytes = f.read()

            st.download_button(label="📄 Download DOCX", 
                               data=docx_bytes, 
                               file_name="generated_assignment.docx", 
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"❌ An error occurred while generating the DOCX: {e}")
